from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)

title = doc.add_heading('银行外汇交易系统菜单目录', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('')

menu_groups = [
    {
        'title': '外汇交易管理',
        'items': [
            {'label': '外汇工作台', 'path': '/fx/workbench'},
            {'label': '未到期交易管理', 'path': '/fx/unmatured'},
            {'label': '客户交易查询', 'path': '/fx/customer-query'},
            {'label': '待办任务', 'path': '/fx/todo'}
        ],
        'sub_items': [
            {'label': '即期交易录入', 'path': '/fx/spot-entry'},
            {'label': '远期交易录入', 'path': '/fx/forward-entry'},
            {'label': '掉期交易录入', 'path': '/fx/swap-entry'},
            {'label': '提前违约', 'path': '/fx/early-default'},
            {'label': '提前交割', 'path': '/fx/early-delivery'},
            {'label': '掉期全部违约', 'path': '/fx/swap-full-default'},
            {'label': '远期展期', 'path': '/fx/rollover'}
        ]
    },
    {
        'title': '期权交易管理',
        'items': [
            {'label': '期权工作台', 'path': '/option/workbench'},
            {'label': '期权交易录入', 'path': '/option/entry'},
            {'label': '期权交易复核', 'path': '/option/review'},
            {'label': '期权存续期管理', 'path': '/option/lifecycle'},
            {'label': '期权交易查询', 'path': '/option/query'}
        ],
        'sub_items': []
    },
    {
        'title': '公共管理',
        'items': [
            {'label': '系统参数管理', 'path': '/system/param'},
            {'label': '客户管理', 'path': '/system/customer'},
            {'label': '登录用户管理', 'path': '/system/user'},
            {'label': '定时任务', 'path': '/system/task'}
        ],
        'sub_items': []
    }
]

for idx, group in enumerate(menu_groups, 1):
    heading = doc.add_heading(f'{idx}. {group["title"]}', level=1)
    heading.style.font.name = '微软雅黑'
    
    for jdx, item in enumerate(group['items'], 1):
        p = doc.add_paragraph()
        p.add_run(f'{idx}.{jdx} {item["label"]}').bold = True
        p.add_run(f'\n路径: {item["path"]}').font.size = Pt(9)
    
    if group['sub_items']:
        sub_heading = doc.add_heading('交易录入', level=2)
        for kdx, sub_item in enumerate(group['sub_items'], 1):
            p = doc.add_paragraph()
            p.add_run(f'{idx}.{kdx} {sub_item["label"]}').bold = True
            p.add_run(f'\n路径: {sub_item["path"]}').font.size = Pt(9)
    
    doc.add_paragraph('')

doc.save('c:/workspace/xfunds-system-master/xfunds-system-master/系统菜单目录.docx')
print('文档已生成: 系统菜单目录.docx')